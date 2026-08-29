"""RAG 文档加载、切分、增量索引和隔离测试。"""

from collections.abc import Sequence
from pathlib import Path

from harness.capabilities.rag import (
    AccessScope,
    DocumentChunk,
    DocumentIndexState,
)
from services.rag import (
    DocumentSplitter,
    IngestionService,
    InMemoryVectorStore,
    TextSplitterConfig,
)
from services.rag.ingestion import load_source_documents


class FakeEmbeddings:
    model_name = "fake-embedding-v1"
    dimension = 3

    def __init__(self) -> None:
        self.document_batches: list[tuple[str, ...]] = []

    def embed_documents(self, texts: Sequence[str]) -> Sequence[Sequence[float]]:
        self.document_batches.append(tuple(texts))
        return tuple(self._vector(text) for text in texts)

    def embed_query(self, text: str) -> Sequence[float]:
        return self._vector(text)

    @staticmethod
    def _vector(text: str) -> tuple[float, float, float]:
        normalized = text.casefold()
        return (
            float("refund" in normalized or "退款" in normalized),
            float("alice" in normalized),
            1.0,
        )


def test_loader_parses_frontmatter_and_splitter_preserves_sections(tmp_path: Path) -> None:
    (tmp_path / "guide.md").write_text(
        "---\ntitle: Product Guide\ncategory: policy\ntags: [refund, order]\n---\n"
        "# Refunds\nApply in seven days.\n\n## Exceptions\nVirtual goods are excluded.",
        encoding="utf-8",
    )
    (tmp_path / "notes.txt").write_text("Plain text notes", encoding="utf-8")

    documents = load_source_documents(
        tmp_path,
        knowledge_base_id="kb",
        scope="public",
    )
    markdown = next(document for document in documents if document.source == "guide.md")
    chunks = DocumentSplitter(TextSplitterConfig(chunk_size=80, chunk_overlap=10)).split(
        markdown
    )

    assert len(documents) == 2
    assert markdown.metadata["category"] == "policy"
    assert markdown.metadata["scope"] == "public"
    assert {str(chunk.metadata["section"]) for chunk in chunks} == {
        "Refunds",
        "Refunds > Exceptions",
    }
    assert all(chunk.text.strip() for chunk in chunks)
    assert all(chunk.metadata["chunk_id"] == chunk.chunk_id for chunk in chunks)
    assert chunks == DocumentSplitter(
        TextSplitterConfig(chunk_size=80, chunk_overlap=10)
    ).split(markdown)


def test_loader_parses_docx_headings_tables_and_core_title(tmp_path: Path) -> None:
    from docx import Document

    path = tmp_path / "guide.docx"
    document = Document()
    document.core_properties.title = "Product Guide"
    document.add_heading("Refunds", level=1)
    document.add_paragraph("Apply in seven days.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Product"
    table.rows[0].cells[1].text = "Refundable"
    document.save(path)

    documents = load_source_documents(
        tmp_path,
        knowledge_base_id="kb",
        scope="public",
    )

    assert len(documents) == 1
    docx_document = documents[0]
    assert docx_document.source == "guide.docx"
    assert docx_document.metadata["title"] == "Product Guide"
    assert "# Refunds" in docx_document.text
    assert "Product | Refundable" in docx_document.text
    chunks = DocumentSplitter(TextSplitterConfig(chunk_size=100, chunk_overlap=10)).split(
        docx_document
    )
    assert {str(chunk.metadata["section"]) for chunk in chunks} == {"Refunds"}


def test_ingestion_skips_unchanged_and_replaces_changed_document(tmp_path: Path) -> None:
    path = tmp_path / "guide.txt"
    path.write_text("refund policy version one", encoding="utf-8")
    embeddings = FakeEmbeddings()
    store = InMemoryVectorStore()
    service = IngestionService(
        embeddings,
        store,
        DocumentSplitter(TextSplitterConfig(chunk_size=100, chunk_overlap=10)),
        knowledge_base_id="kb",
    )

    first = service.index_directory(tmp_path, scope="public")
    second = service.index_directory(tmp_path, scope="public")
    path.write_text("refund policy version two", encoding="utf-8")
    third = service.index_directory(tmp_path, scope="public")

    assert first.indexed == 1 and not first.failed
    assert second.skipped == 1 and len(embeddings.document_batches) == 2
    assert third.indexed == 1 and third.deleted_chunks == 1
    hits = store.search((1.0, 0.0, 1.0), AccessScope(), 5, {})
    assert len(hits) == 1
    assert "version two" in hits[0].chunk.text


def test_ingestion_reindexes_frontmatter_changes(tmp_path: Path) -> None:
    path = tmp_path / "guide.md"
    path.write_text(
        "---\ntitle: Guide\ncategory: policy-v1\n---\nRefund policy",
        encoding="utf-8",
    )
    embeddings = FakeEmbeddings()
    store = InMemoryVectorStore()
    service = IngestionService(
        embeddings,
        store,
        DocumentSplitter(TextSplitterConfig(chunk_size=100, chunk_overlap=10)),
        knowledge_base_id="kb",
    )

    first = service.index_directory(tmp_path, scope="public")
    path.write_text(
        "---\ntitle: Guide\ncategory: policy-v2\n---\nRefund policy",
        encoding="utf-8",
    )
    second = service.index_directory(tmp_path, scope="public")

    assert first.indexed == 1
    assert second.indexed == 1
    assert second.skipped == 0
    assert second.deleted_chunks == 1
    document_id = load_source_documents(
        tmp_path,
        knowledge_base_id="kb",
        scope="public",
    )[0].document_id
    hits = store.search((1.0, 0.0, 1.0), AccessScope(), 5, {})
    assert len(hits) == 1
    assert hits[0].chunk.document_id == document_id
    assert hits[0].chunk.metadata["category"] == "policy-v2"


def test_vector_store_enforces_public_and_current_user_scope() -> None:
    store = InMemoryVectorStore()
    chunks = (
        DocumentChunk(
            document_id="public-doc",
            chunk_id="public",
            text="public",
            metadata={"scope": "public", "source": "public.md", "section": "Public"},
        ),
        DocumentChunk(
            document_id="alice-doc",
            chunk_id="alice",
            text="alice",
            metadata={
                "scope": "user",
                "user_id": "alice",
                "source": "alice.md",
                "section": "Alice",
            },
        ),
        DocumentChunk(
            document_id="bob-doc",
            chunk_id="bob",
            text="bob",
            metadata={
                "scope": "user",
                "user_id": "bob",
                "source": "bob.md",
                "section": "Bob",
            },
        ),
    )
    for chunk in chunks:
        store.replace_document(
            DocumentIndexState(
                document_id=chunk.document_id,
                content_hash=chunk.chunk_id,
                embedding_model="fake",
                embedding_dimension=2,
                splitter_version="v1",
                chunk_ids=(chunk.chunk_id,),
            ),
            (chunk,),
            ((1.0, 0.0),),
        )

    hits = store.search((1.0, 0.0), AccessScope(user_id="alice"), 10, {})

    assert {hit.chunk.chunk_id for hit in hits} == {"public", "alice"}
    assert "bob" not in {hit.chunk.chunk_id for hit in hits}
